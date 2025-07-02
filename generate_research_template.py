from docx import Document
from docx.shared import Inches

class ResearchTemplate:
    def __init__(self,level,topic,total_pictures,total_videos,questions,folder_name):
        self.level=level
        self.topic=topic
        self.total_pictures=total_pictures
        self.total_videos=total_videos
        self.questions=questions
        self.folder_name=folder_name
    
    def add_list_style(self,list_paragraphs,num_id_list_number_new):
        for i in range(len(list_paragraphs)):
            numPr = list_paragraphs[i]._element.pPr._add_numPr()
            numPr._add_numId().val = num_id_list_number_new

    def add_list_paragraph(self,research_document,list_paragraphs,text):
        list_paragraphs.append(research_document.add_paragraph(text,style="List Number"))

    def generate_new_list_style(self,document):
        #prepare the numberings to have a new numbering, which points to the same abstract numbering, 
        #the style "List Number" also points to but has start override set
        styles = document.styles
        #get numId to which style 'List Number' links
        num_id_list_number = -1
        for style in styles:
            if (style.name == 'List Number'):
                num_id_list_number = style._element.pPr.numPr.numId.val
        #add new numbering linking to same abstractNumId but has startOverride 
        #and get new numId
        num_id_list_number_new = -1
        if (num_id_list_number > -1):        
            ct_numbering = document.part.numbering_part.numbering_definitions._numbering
            ct_num = ct_numbering.num_having_numId(num_id_list_number)
            abstractNumId = ct_num.abstractNumId.val
            ct_num = ct_numbering.add_num(abstractNumId)
            num_id_list_number_new = ct_num.numId
            startOverride = ct_num.add_lvlOverride(0)._add_startOverride()
            startOverride.val = 1
        return num_id_list_number_new

    def generate_research_documents(self):
        #Initial Setup
        
        research_document = Document()
        research_document.add_heading(f'{self.topic} - {self.level}', 0)
        research_document.add_heading(f'Today you are going to do some research on {self.topic}!', 1)
        self.write_research(research_document)
        self.write_pictures(research_document)
        self.write_videos(research_document)
        
        research_document.save(f'{self.folder_name}/Part 1 - {self.topic} {self.level} Research Template.docx')

class LLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        research_document.add_heading(f'Writing Goals',2)
        research_document.add_paragraph("Low Level is an easier Soft Skills research level. For this level, you are given tables. You have to find links to answer questions. Repeated steps are written out.\n\nAvoid using information directly from Generative AI (Copilot Answer, AI Overview, ChatGPT, Google Gemini). Make sure to use information from a website and include the link to that website. Try to avoid copying and pasting too much information. Only use necessary information from websites.")

        #Part 1 Research
        research_document.add_heading(f'Part 1 - Research',2)
        research_document.add_paragraph('Please fill in the chart below using google to help you find the answers. Attach the link to the website you used')
        research_document.add_paragraph('Make sure to put information in your own words. Do not copy and paste information directly.')

        table = research_document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Answer'
        hdr_cells[2].text = 'Website Link'
        for i in range(len(self.questions)):
            row_cells = table.add_row().cells
            row_cells[0].text = self.questions[i]
    
    def write_pictures(self,research_document):
        #Part 2 Pictures
        research_document.add_heading(f'Part 2 - Pictures',2)
        p = research_document.add_paragraph(f'Research and find {self.total_pictures} pictures for {self.topic}. In google images type “{self.topic} Tips.” Click on the picture and copy and paste it into the table below. Attach the link to the website you used below.')
        table = research_document.add_table(rows=1+self.total_pictures, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Pictures'
        hdr_cells[1].text = 'Website'
    
    def write_videos(self,research_document):
        #Part 3 Videos
        research_document.add_heading(f'Part 3 - Videos',2)
        p = research_document.add_paragraph(f'Please watch and find {self.total_videos} YouTube videos about {self.topic} tips for students. These should be videos YOU find fun and interesting. Copy and paste the links below. In the comments column write 1-2 sentences about your favorite part of the video.')

        table = research_document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Video Links'
        hdr_cells[1].text = 'Comments'
        for i in range(self.total_videos):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{i+1}."

class MLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        research_document.add_heading(f'Writing Goals',2)
        research_document.add_paragraph("Medium Level is the average Soft Skills research level. For this level, you have to build tables and find links to answer questions. Repeated steps are only written out once.\n\nAvoid using information directly from Generative AI (Copilot Answer, AI Overview, ChatGPT, Google Gemini). Make sure to use information from a website and include the link to that website. Avoid copying and pasting information. Try to rewrite information in your own words.")

        research_document.add_heading(f'Part 1 - Research',2)
        research_document.add_paragraph("INSERT TABLE HERE")
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]

        self.add_list_paragraph(research_document,list_paragraphs,"Click above where it says 'INSERT TABLE HERE'. Click Insert, then click on table.")
        self.add_list_paragraph(research_document,list_paragraphs,"Use your mouse to highlight a 3x6 area.")
        research_document.add_picture("images/insert_table.png")
        research_document.add_picture("images/table_3x6.png")
        self.add_list_paragraph(research_document,list_paragraphs,"Once you have a 3x6 click on the box, a table should pop up in your document.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Question")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Answer")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 3 Website Link")     
        for i in range(len(self.questions)):
            self.add_list_paragraph(research_document,list_paragraphs,f"In column 1, row {i+2}, type the question {self.questions[i]}")
        #list_number+=1
        
        self.add_list_style(list_paragraphs,num_id_list_number_new)


        research_document.add_heading(f'Part 2 - Research',2)
        research_document.add_paragraph("Please fill in the chart you made above using google to help you find the answers. Type each question into google and find an answer. Type your answers in the box next to the question. Do not copy and paste your answers directly. Then, copy and paste the link into the third box. Attach the link to the website you used to find the answer.")

    def write_pictures(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]
        
        research_document.add_heading(f'Part 3 - Pictures',2)
        research_document.add_paragraph("INSERT TABLE HERE")
        self.add_list_paragraph(research_document,list_paragraphs,"Click above where it says 'INSERT TABLE HERE'. Click Insert, then click on table.")
        self.add_list_paragraph(research_document,list_paragraphs,"Use your mouse to highlight a 2x6 area.")
        research_document.add_picture("images/insert_table.png")
        research_document.add_picture("images/table_2x6.png")
        self.add_list_paragraph(research_document,list_paragraphs,"Once you have a 2x6 click on the box, a table should pop up in your document.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Pictures")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Website")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 4 - Pictures',2)
        research_document.add_paragraph(f"Research and find 5 pictures for {self.topic}. In google images type {self.topic} Click on the picture and copy and paste it into the table you made above. Attach the link to the website you used in the chart above.")

    def write_videos(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]
    
        research_document.add_heading(f'Part 5 - Videos',2)
        research_document.add_paragraph("INSERT TABLE HERE")
        self.add_list_paragraph(research_document,list_paragraphs,"Click above where it says 'INSERT TABLE HERE'. Click Insert, then click on table.")
        self.add_list_paragraph(research_document,list_paragraphs,"Use your mouse to highlight a 2x6 area.")
        research_document.add_picture("images/insert_table.png")
        research_document.add_picture("images/table_2x6.png")
        self.add_list_paragraph(research_document,list_paragraphs,"Once you have a 2x6 click on the box, a table should pop up in your document.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Video Links")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Comments")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 6 - Videos',2)
        research_document.add_paragraph(f"Please find and watch 5 YouTube videos about {self.topic} tips for students. These should be videos YOU find fun and interesting. Copy and paste the links in your table above. When you have finished watching a video write 1-2 sentences about what you liked best about the video in the box next to where you put the link.")

class HLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        research_document.add_heading(f'Writing Goals',2)
        research_document.add_paragraph("High Level is the hardest regular Soft Skills research level. For this level, you have to build tables and find links to answer questions. Repeated steps are only written out once.\n\nDo not use information directly from Generative AI (Copilot Answer, AI Overview, ChatGPT, Google Gemini). Make sure to use information from a website and include the link to that website. Do not copy and paste information. Rewrite information in your own words.")

        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]
        research_document.add_heading(f'Part 1 - Research',2)

        self.add_list_paragraph(research_document,list_paragraphs,"Create a table with 3 columns and 6 rows.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Question")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Answer")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 3 Website Link")
        for i in range(len(self.questions)):
           self.add_list_paragraph(research_document,list_paragraphs,f"In column 1, row {i+2}, type the question {self.questions[i]}")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 2 - Research',2)
        research_document.add_paragraph("Please fill in the chart you made. Use Google to help you find the answers. Type all your answers. Do not copy and paste directly. Attach the link to the website you used. You may copy and paste the link.")

    def write_pictures(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]

        research_document.add_heading(f'Part 3 - Pictures',2)
        self.add_list_paragraph(research_document,list_paragraphs,"Create a table with 2 columns and 6 rows.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Pictures")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Website")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 4 - Pictures',2)
        research_document.add_paragraph(f"Research and find 5 pictures for {self.topic}. In google images type {self.topic} Click on the picture and copy and paste it into the table you made above. Attach the link to the website you used in the chart above.")

    def write_videos(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]

        research_document.add_heading(f'Part 5 - Videos',2)
        self.add_list_paragraph(research_document,list_paragraphs,"Create a table with 2 columns and 6 rows.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Video Links")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Comments")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 6 - Videos',2)
        research_document.add_paragraph(f"Please find and watch 5 YouTube videos about {self.topic} tips for students. These should be videos YOU find fun and interesting. Copy and paste the links in your table above. When you have finished watching a video write 1-2 sentences about what you liked best about the video in the box next to where you put the link.")

class VHLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]
        research_document.add_heading(f'Assignment Writing Goals',2)
        research_document.add_paragraph("Very High Level is a special Soft Skills research level. This is significantly harder than other levels. For this level, you have to determine what questions you want to answer. Choose interesting questions that will help your audience understand your topic.\n\nDo not copy and paste information. Do not use information directly from Generative AI (Copilot Answer, AI Overview, ChatGPT, Google Gemini). Make sure to use information from a website and include the link to that website. You want to make a clear and concise presentation. Good luck.")

        research_document.add_heading(f'Writing Advice',2)
        self.add_list_paragraph(research_document,list_paragraphs,f"Read this section carefully before you start. This assignment is more difficult, but more fuffilling than other Soft Skills Research assignments.") 
        self.add_list_paragraph(research_document,list_paragraphs,f"For this assignment, you have much more freedom in what you will talk about for {self.topic}.") 
        self.add_list_paragraph(research_document,list_paragraphs,f"You are going to determine four of the questions. Your goal is put together a strong presentation with more of your own ideas.") 
        self.add_list_paragraph(research_document,list_paragraphs,f"This will be tougher but I think you can do it.") 
        self.add_list_paragraph(research_document,list_paragraphs,f"Make sure to paraphrase information you find. Do not copy and paste information directly. Write information in your own words.") 
        self.add_list_paragraph(research_document,list_paragraphs,f"You want to rewrite other people's ideas to match how you would say those ideas.")
        self.add_list_paragraph(research_document,list_paragraphs,f"You would be speaking when presenting this, so avoid using too many words.")
        self.add_list_paragraph(research_document,list_paragraphs,"You are delivering this presentation to high schoolers, so make sure vocabulary and content are appropriate for high schoolers. Avoid overly simple or overly complicated information.")
        self.add_list_paragraph(research_document,list_paragraphs,f"Try to make your questions creative and interesting. Avoid listing for answers and try to provide explanations for ideas.")
        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 0 - Your Main Idea',2)
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]
        self.add_list_paragraph(research_document,list_paragraphs,f"Before you start researching about {self.topic} you want to determine what you will research.")
        self.add_list_paragraph(research_document,list_paragraphs,f"This is your main idea.")
        self.add_list_paragraph(research_document,list_paragraphs,f"Your main idea should be a sentence or the title of your presentation.")
        self.add_list_paragraph(research_document,list_paragraphs,f"Write down your main idea below.")
        self.add_list_style(list_paragraphs,num_id_list_number_new)

        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]
        research_document.add_heading(f'Part 1 - Research',2)
        self.add_list_paragraph(research_document,list_paragraphs,"Create a table with 3 columns and 6 rows.")
        self.add_list_paragraph(research_document,list_paragraphs,"Look up how to make a Microsoft Word table if you do not know how to do so.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Question")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Answer")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 3 Website Link")
        
        self.add_list_paragraph(research_document,list_paragraphs,f"In column 1, row 2, type the question {self.questions[0]}")
        self.add_list_paragraph(research_document,list_paragraphs,f"Add four more questions. You can decide what questions you want to answer. Your answers should give your audience a general idea about {self.topic}.")

        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 2 - Research',2)
        research_document.add_paragraph("Please fill in the chart you made. Research your answers online.") 
        research_document.add_paragraph("Make sure to include the link to the website.")

    def write_pictures(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]

        research_document.add_heading(f'Part 3 - Pictures',2)
        self.add_list_paragraph(research_document,list_paragraphs,"Create a table with 2 columns and 6 rows.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Pictures")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Website")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 4 - Pictures',2)
        research_document.add_paragraph(f"Research and find 5 pictures for {self.topic}. Make sure the pictures relate back to your answers.")
        research_document.add_paragraph(f"Add the pictures and the picture links to your table.")

    def write_videos(self,research_document):
        num_id_list_number_new=self.generate_new_list_style(research_document)
        list_paragraphs=[]

        research_document.add_heading(f'Part 5 - Videos',2)
        self.add_list_paragraph(research_document,list_paragraphs,"Create a table with 2 columns and 6 rows.")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 1 Video Links")
        self.add_list_paragraph(research_document,list_paragraphs,"Label the first box in column 2 Comments")
        #list_number+=1

        self.add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 6 - Videos',2)
        research_document.add_paragraph(f"Please find and watch 5 YouTube videos about {self.topic}. These should be videos YOU find fun and interesting. When you have finished watching a video write 1-2 sentences about what you liked best about the video in the box next to where you put the link.")
        
        research_document.add_paragraph(f"Add the video links and your comments to your table.")
        research_document.add_paragraph(f"Make sure the videos relate back to your answers.")