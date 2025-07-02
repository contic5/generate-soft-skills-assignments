from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_numbering_style(doc, num_id):
    # Create a numbering definition for the list
    numbering = doc.part.numbering_part.element

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(num_id))

    multi_level = OxmlElement('w:multiLevelType')
    multi_level.set(qn('w:val'), 'singleLevel')
    abstract_num.append(multi_level)

    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')

    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)

    num_fmt = OxmlElement('w:numFmt')
    num_fmt.set(qn('w:val'), 'decimal')
    lvl.append(num_fmt)

    lvl_text = OxmlElement('w:lvlText')
    lvl_text.set(qn('w:val'), '%1.')
    lvl.append(lvl_text)

    lvl_jc = OxmlElement('w:lvlJc')
    lvl_jc.set(qn('w:val'), 'left')
    lvl.append(lvl_jc)

    p_pr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    p_pr.append(ind)
    lvl.append(p_pr)

    abstract_num.append(lvl)
    numbering.append(abstract_num)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))

    abstract_num_id_elem = OxmlElement('w:abstractNumId')
    abstract_num_id_elem.set(qn('w:val'), str(num_id))
    num.append(abstract_num_id_elem)
    numbering.append(num)

def add_numbered_list(doc, items, num_id):
    create_numbering_style(doc, num_id)
    
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)

        num_pr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        num_id_elem = OxmlElement('w:numId')
        num_id_elem.set(qn('w:val'), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_elem)

        p._p.get_or_add_pPr().append(num_pr)

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Multiple Numbered Lists with Restarted Numbering', level=1)

# First numbered list
list1 = ["Item 1.1", "Item 1.2", "Item 1.3"]
add_numbered_list(doc, list1, num_id=1)

# Blank line to separate lists
doc.add_paragraph()

# Second numbered list
list2 = ["Item 2.1", "Item 2.2", "Item 2.3"]
add_numbered_list(doc, list2, num_id=2)

# Blank line to separate lists
doc.add_paragraph()

# Third numbered list
list3 = ["Item 3.1", "Item 3.2", "Item 3.3"]
add_numbered_list(doc, list3, num_id=3)

# Save the document
doc.save('multiple_numbered_lists_restart.docx')
