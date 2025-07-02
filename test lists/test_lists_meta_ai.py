from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

document = Document()

# First list
for i in range(1, 4):
    paragraph = document.add_paragraph(f'List 1 item {i}')
    paragraph.style = document.styles['List Number']

# Restart list
paragraph = document.add_paragraph('restart')
paragraph.style = document.styles['List Number']
paragraph.paragraph_format.start_indent = Inches(0.5)
paragraph.paragraph_format.space_before = Inches(0)
paragraph.paragraph_format.space_after = Inches(0)
paragraph.paragraph_format.first_line_indent = Inches(-0.5)
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph.paragraph_format.page_break_before = True

# Hide the "restart" paragraph
paragraph.font.size = Pt(1)
paragraph.font.hidden = True

# Second list, starting from 1
section = document.add_section(WD_SECTION.NEW_PAGE)
for i in range(1, 4):
    paragraph = document.add_paragraph(f'List 2 item {i}')
    paragraph.style = document.styles['List Number']

# Restart list
paragraph = document.add_paragraph('restart')
paragraph.style = document.styles['List Number']
paragraph.paragraph_format.start_indent = Inches(0.5)
paragraph.paragraph_format.space_before = Inches(0)
paragraph.paragraph_format.space_after = Inches(0)
paragraph.paragraph_format.first_line_indent = Inches(-0.5)
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph.paragraph_format.page_break_before = True

# Hide the "restart" paragraph
paragraph.font.size = Pt(1)
paragraph.font.hidden = True

# Third list, starting from 1
section = document.add_section(WD_SECTION.NEW_PAGE)
for i in range(1, 4):
    paragraph = document.add_paragraph(f'List 3 item {i}')
    paragraph.style = document.styles['List Number']

document.save('lists.docx')