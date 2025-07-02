from docx import Document

document = Document()

paragraph = document.add_paragraph()
paragraph.add_run('List 1').bold = True

for i in range(1, 4):
    document.add_paragraph(
        f'List 1 Item {i}', style='List Number'
    )
    
paragraph = document.add_paragraph()
paragraph.add_run('List 2').bold = True

#prepare the numberings to have a new numbering, which points to the same abstract numbering, 
#the style "List Number" also points to but has start override set
styles = document.styles
#get numId to which style 'List Number' links
num_id_list_number = -1
for style in styles:
    if (style.name == 'List Number'):
        num_id_list_number = style._element.pPr.numPr.numId.val
#add new numbering linking to same abstractNumId but has startOverride 
#and get new numId
num_id_list_number_new = -1
if (num_id_list_number > -1):        
    ct_numbering = document.part.numbering_part.numbering_definitions._numbering
    ct_num = ct_numbering.num_having_numId(num_id_list_number)
    abstractNumId = ct_num.abstractNumId.val
    ct_num = ct_numbering.add_num(abstractNumId)
    num_id_list_number_new = ct_num.numId
    startOverride = ct_num.add_lvlOverride(0)._add_startOverride()
    startOverride.val = 1

for i in range(1, 4):
    paragraph = document.add_paragraph(
            f'List 2 Item {i}', style='List Number'
        )
    
    #first paragraph in new list links to new numId having startOverride
    if ( i == 1 and num_id_list_number_new > -1):
        numPr = paragraph._element.pPr._add_numPr()
        numPr._add_numId().val = num_id_list_number_new

document.save('test.docx')